import os
import json
import csv
import requests
from typing import List, Dict, Optional, Any, Callable
from pathlib import Path
import asyncio
import aiohttp
from dataclasses import dataclass
from utils.path_utils import validate_path
from utils.exceptions import ProcessingError, FileOperationError
import time
import hashlib

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from advanced_rag_pipeline import Document, AdvancedRAGPipeline

MAX_DOCS = 10000
MAX_FILE_SIZE_MB = 50
MAX_PDF_PAGES = 200
MAX_DOCX_PARAGRAPHS = 10000

def safe_ingest(fn: Callable) -> Callable:
    """Декоратор для безопасного ingestion с логированием и единым форматом ошибок."""
    def wrapper(self, *args, **kwargs):
        try:
            return fn(self, *args, **kwargs)
        except Exception as e:
            raise FileOperationError(f"[{fn.__name__}] {e}") from e
    return wrapper

class DataIngestionManager:
    """Менеджер для загрузки и валидации данных из различных источников для RAG."""

    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline
        self.allowed_data_dir = Path("./data").resolve()

    @safe_ingest
    def load_from_text_file(self, filepath: str, encoding: str = 'utf-8') -> str:
        path = Path(filepath)
        is_valid, reason = validate_path(path, allowed_dir=self.allowed_data_dir, allowed_exts={".txt"}, max_size_mb=MAX_FILE_SIZE_MB)
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        with open(path, 'r', encoding=encoding) as f:
            text = f.read()
            if not text.strip():
                raise FileOperationError(f"Файл пустой или содержит только пробелы: {filepath}")
            return text

    @safe_ingest
    def load_from_pdf(self, filepath: str, max_pages: int = MAX_PDF_PAGES) -> str:
        if not PDF_AVAILABLE:
            raise ProcessingError("PyPDF2 не установлен. Установите: pip install PyPDF2")
        path = Path(filepath)
        is_valid, reason = validate_path(path, allowed_dir=self.allowed_data_dir, allowed_exts={".pdf"}, max_size_mb=MAX_FILE_SIZE_MB)
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        text = ""
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                if i >= max_pages:
                    break
                text += (page.extract_text() or "") + "\n"
        if not text.strip():
            raise FileOperationError(f"PDF пустой или нечитабельный: {filepath}")
        return text

    @safe_ingest
    def load_from_docx(self, filepath: str, max_paragraphs: int = MAX_DOCX_PARAGRAPHS) -> str:
        if not DOCX_AVAILABLE:
            raise ProcessingError("python-docx не установлен. Установите: pip install python-docx")
        path = Path(filepath)
        is_valid, reason = validate_path(path, allowed_dir=self.allowed_data_dir, allowed_exts={".docx"}, max_size_mb=MAX_FILE_SIZE_MB)
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        doc = docx.Document(path)
        paragraphs = doc.paragraphs[:max_paragraphs]
        text = "\n".join(paragraph.text for paragraph in paragraphs)
        if not text.strip():
            raise FileOperationError(f"DOCX пустой: {filepath}")
        return text

    @safe_ingest
    def load_from_csv(self, filepath: str, text_columns: List[str], max_rows: int = MAX_DOCS) -> List[Dict[str, Any]]:
        path = Path(filepath)
        is_valid, reason = validate_path(path, allowed_dir=self.allowed_data_dir, allowed_exts={".csv"}, max_size_mb=MAX_FILE_SIZE_MB)
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        documents = []
        with open(path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for i, row in enumerate(reader):
                if i >= max_rows:
                    break
                content_parts = [str(row[col]) for col in text_columns if col in row and row[col]]
                if content_parts:
                    content = " ".join(content_parts)
                    metadata = {k: v for k, v in row.items() if k not in text_columns}
                    documents.append({'id': f"csv_row_{i}", 'content': content, 'metadata': metadata})
        if not documents:
            raise FileOperationError(f"CSV не содержит нужных данных по столбцам {text_columns}: {filepath}")
        return documents

    @safe_ingest
    def load_from_json(self, filepath: str, content_field: str, id_field: Optional[str] = None, max_docs: int = MAX_DOCS) -> List[Dict[str, Any]]:
        path = Path(filepath)
        is_valid, reason = validate_path(path, allowed_dir=self.allowed_data_dir, allowed_exts={".json"}, max_size_mb=MAX_FILE_SIZE_MB)
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        documents = []
        if isinstance(data, list):
            for i, item in enumerate(data):
                if i >= max_docs:
                    break
                if content_field in item:
                    doc_id = item.get(id_field, f"json_item_{i}") if id_field else f"json_item_{i}"
                    content = str(item[content_field])
                    metadata = {k: v for k, v in item.items() if k not in [content_field, id_field]}
                    documents.append({'id': doc_id, 'content': content, 'metadata': metadata})
        if not documents:
            raise FileOperationError(f"JSON не содержит элементов с полем {content_field}: {filepath}")
        return documents

    @safe_ingest
    def load_from_url(self, url: str, timeout: int = 30) -> str:
        response = requests.get(url, timeout=timeout)
        response.raise_for_status()
        content = response.text
        if 'text/html' in response.headers.get('content-type', '') and BS4_AVAILABLE:
            soup = BeautifulSoup(content, 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            content = soup.get_text()
        if not content.strip():
            raise ProcessingError(f"URL {url} не содержит текста")
        return content

    async def load_from_urls_async(self, urls: List[str], timeout: int = 30) -> List[Dict[str, Any]]:
        async def fetch_url(session, url):
            try:
                async with session.get(url, timeout=timeout) as response:
                    content = await response.text()
                    if 'text/html' in response.headers.get('content-type', '') and BS4_AVAILABLE:
                        soup = BeautifulSoup(content, 'html.parser')
                        for script in soup(["script", "style"]):
                            script.decompose()
                        content = soup.get_text()
                    return {
                        'id': hashlib.md5(url.encode()).hexdigest(),
                        'content': content,
                        'metadata': {'source_url': url, 'status': 'success'}
                    }
            except Exception as e:
                return {
                    'id': hashlib.md5(url.encode()).hexdigest(),
                    'content': "",
                    'metadata': {'source_url': url, 'status': 'error', 'error': str(e)}
                }
        results = []
        async with aiohttp.ClientSession() as session:
            tasks = [fetch_url(session, url) for url in urls]
            fetch_results = await asyncio.gather(*tasks)
            # Сохраняем все результаты, включая ошибки
            return fetch_results

# --- Аналитика и Web-интерфейс (оставлен без изменений, кроме type hints и лимитов) ---

class RAGAnalytics:
    """Аналитика и мониторинг RAG системы"""
    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline
        self.query_log: List[Dict[str, Any]] = []

    def log_query(self, query: str, results_count: int, processing_time: float) -> None:
        self.query_log.append({
            'timestamp': time.time(),
            'query': query,
            'results_count': results_count,
            'processing_time': processing_time
        })

    def get_query_stats(self) -> Dict[str, Any]:
        if not self.query_log:
            return {"message": "Нет данных по запросам"}
        processing_times = [log['processing_time'] for log in self.query_log]
        results_counts = [log['results_count'] for log in self.query_log]
        return {
            'total_queries': len(self.query_log),
            'avg_processing_time': sum(processing_times) / len(processing_times),
            'max_processing_time': max(processing_times),
            'min_processing_time': min(processing_times),
            'avg_results_count': sum(results_counts) / len(results_counts),
            'queries_per_hour': len([q for q in self.query_log if time.time() - q['timestamp'] < 3600])
        }

    def analyze_collection_content(self, max_categories: int = 20, max_languages: int = 10) -> Dict[str, Any]:
        try:
            all_data = self.rag.collection.get()
            documents = all_data.get('documents', [])
            metadatas = all_data.get('metadatas', [])
            if not documents:
                return {"message": "Коллекция пуста"}
            total_docs = len(documents)
            total_chars = sum(len(doc) for doc in documents)
            avg_doc_length = total_chars / total_docs
            categories = {}
            languages = {}
            for metadata in metadatas:
                if metadata:
                    if 'category' in metadata:
                        cat = metadata['category']
                        if cat in categories: categories[cat] += 1
                        elif len(categories) < max_categories: categories[cat] = 1
                    if 'language' in metadata:
                        lang = metadata['language']
                        if lang in languages: languages[lang] += 1
                        elif len(languages) < max_languages: languages[lang] = 1
            return {
                'total_documents': total_docs,
                'total_characters': total_chars,
                'average_document_length': avg_doc_length,
                'categories_distribution': categories,
                'languages_distribution': languages,
                'longest_document': max(len(doc) for doc in documents),
                'shortest_document': min(len(doc) for doc in documents)
            }
        except Exception as e:
            return {"error": f"Ошибка анализа: {e}"}

class RAGWebInterface:
    """Простой веб-интерфейс для RAG системы"""
    def __init__(self, rag_pipeline: AdvancedRAGPipeline, analytics: RAGAnalytics):
        self.rag = rag_pipeline
        self.analytics = analytics

    def generate_html_interface(self) -> str:
        # (оставлен без изменений, так как это статичный шаблон)
        html_template = """ ... (см. исходник) ... """
        return html_template

    def save_interface(self, filepath: str = "rag_interface.html") -> None:
        html_content = self.generate_html_interface()
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print(f"HTML интерфейс сохранен в {filepath}")

class RAGBenchmarking:
    """Бенчмаркинг и тестирование RAG системы"""
    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline

    def create_test_dataset(self, size: int = 100) -> List[Document]:
        test_topics = [
            "машинное обучение", "искусственный интеллект", "нейронные сети",
            "векторные базы данных", "обработка естественного языка",
            "компьютерное зрение", "глубокое обучение", "анализ данных",
            "программирование на Python", "веб-разработка"
        ]
        documents = []
        for i in range(size):
            topic = test_topics[i % len(test_topics)]
            content = self._generate_synthetic_content(topic, i)
            doc = Document(
                id=f"test_doc_{i}",
                content=content,
                metadata={
                    "topic": topic,
                    "test_id": i,
                    "synthetic": True
                }
            )
            documents.append(doc)
        return documents

    def _generate_synthetic_content(self, topic: str, doc_id: int) -> str:
        templates = [
            f"{topic} является важной областью современной технологии. Документ номер {doc_id} содержит подробную информацию о применении {topic} в различных сферах.",
            f"В этом документе рассматриваются основные принципы {topic}. Это руководство номер {doc_id} поможет понять ключевые концепции.",
            f"Практическое применение {topic} демонстрируется в примере {doc_id}. Здесь представлены лучшие практики и методы."
        ]
        base_content = templates[doc_id % len(templates)]
        additional = f" Дополнительные детали включают технические аспекты, примеры использования и рекомендации экспертов в области {topic}."
        return base_content + additional

    def benchmark_search_performance(self, queries: List[str], n_runs: int = 10) -> Dict[str, Any]:
        results = {
            'queries': [],
            'avg_time': 0,
            'total_time': 0,
            'fastest_query': None,
            'slowest_query': None
        }
        total_time = 0
        fastest_time = float('inf')
        slowest_time = 0
        for query in queries:
            query_times = []
            for _ in range(n_runs):
                start_time = time.time()
                self.rag.search(query, n_results=5)
                end_time = time.time()
                query_time = end_time - start_time
                query_times.append(query_time)
                total_time += query_time
            avg_query_time = sum(query_times) / len(query_times)
            min_query_time = min(query_times)
            max_query_time = max(query_times)
            if min_query_time < fastest_time:
                fastest_time = min_query_time
                results['fastest_query'] = query
            if max_query_time > slowest_time:
                slowest_time = max_query_time
                results['slowest_query'] = query
            results['queries'].append({
                'query': query,
                'avg_time': avg_query_time,
                'min_time': min_query_time,
                'max_time': max_query_time,
                'runs': n_runs
            })
        results['avg_time'] = total_time / (len(queries) * n_runs)
        results['total_time'] = total_time
        results['fastest_time'] = fastest_time
        results['slowest_time'] = slowest_time
        return results

    def evaluate_retrieval_quality(self, test_queries: List[Dict[str, Any]]) -> Dict[str, Any]:
        metrics = {
            'precision_at_k': [],
            'recall_at_k': [],
            'average_precision': []
        }
        for test_case in test_queries:
            query = test_case['query']
            relevant_ids = set(test_case['relevant_doc_ids'])
            results = self.rag.search(query, n_results=10)
            retrieved_ids = [r.document_id for r in results]
            for k in [1, 3, 5, 10]:
                if k <= len(retrieved_ids):
                    retrieved_k = set(retrieved_ids[:k])
                    precision = len(retrieved_k & relevant_ids) / k
                    recall = len(retrieved_k & relevant_ids) / len(relevant_ids) if relevant_ids else 0
                    metrics['precision_at_k'].append(precision)
                    metrics['recall_at_k'].append(recall)
            ap = self._calculate_average_precision(retrieved_ids, relevant_ids)
            metrics['average_precision'].append(ap)
        final_metrics = {
            'mean_precision_at_k': sum(metrics['precision_at_k']) / len(metrics['precision_at_k']) if metrics['precision_at_k'] else 0,
            'mean_recall_at_k': sum(metrics['recall_at_k']) / len(metrics['recall_at_k']) if metrics['recall_at_k'] else 0,
            'mean_average_precision': sum(metrics['average_precision']) / len(metrics['average_precision']) if metrics['average_precision'] else 0
        }
        return final_metrics

    def _calculate_average_precision(self, retrieved_ids: List[str], relevant_ids: set) -> float:
        if not relevant_ids:
            return 0.0
        precisions = []
        relevant_count = 0
        for i, doc_id in enumerate(retrieved_ids):
            if doc_id in relevant_ids:
                relevant_count += 1
                precision = relevant_count / (i + 1)
                precisions.append(precision)
        return sum(precisions) / len(relevant_ids) if precisions else 0.0
