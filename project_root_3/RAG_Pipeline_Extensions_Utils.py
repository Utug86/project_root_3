import os
import json
import csv
import requests
from typing import List, Dict, Optional
from pathlib import Path
import asyncio
import aiohttp
from dataclasses import dataclass
from utils.path_utils import validate_path
from utils.exceptions import ProcessingError, FileOperationError
import time
import hashlib

try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from advanced_rag_pipeline import Document, AdvancedRAGPipeline

class DataIngestionManager:
    """Менеджер для загрузки данных из различных источников"""

    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline
        self.allowed_data_dir = Path("./data").resolve()

    def load_from_text_file(self, filepath: str, encoding: str = 'utf-8') -> str:
        path = Path(filepath)
        is_valid, reason = validate_path(
            path, allowed_dir=self.allowed_data_dir, allowed_exts={".txt"}, max_size_mb=50
        )
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        try:
            with open(path, 'r', encoding=encoding) as f:
                return f.read()
        except Exception as e:
            raise FileOperationError(f"Ошибка чтения файла {filepath}: {e}") from e

    def load_from_pdf(self, filepath: str) -> str:
        if not PDF_AVAILABLE:
            raise ProcessingError("pypdf не установлен. Установите: pip install pypdf")
        path = Path(filepath)
        is_valid, reason = validate_path(
            path, allowed_dir=self.allowed_data_dir, allowed_exts={".pdf"}, max_size_mb=50
        )
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        try:
            text = ""
            with open(path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    text += (page.extract_text() or "") + "\n"
            return text
        except Exception as e:
            raise FileOperationError(f"Ошибка чтения PDF {filepath}: {e}") from e

    def load_from_docx(self, filepath: str) -> str:
        if not DOCX_AVAILABLE:
            raise ProcessingError("python-docx не установлен. Установите: pip install python-docx")
        path = Path(filepath)
        is_valid, reason = validate_path(
            path, allowed_dir=self.allowed_data_dir, allowed_exts={".docx"}, max_size_mb=50
        )
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        try:
            doc = docx.Document(path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            return text
        except Exception as e:
            raise FileOperationError(f"Ошибка чтения DOCX {filepath}: {e}") from e

    def load_from_csv(self, filepath: str, text_columns: List[str]) -> List[Dict]:
        path = Path(filepath)
        is_valid, reason = validate_path(
            path, allowed_dir=self.allowed_data_dir, allowed_exts={".csv"}, max_size_mb=50
        )
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        try:
            documents = []
            with open(path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for i, row in enumerate(reader):
                    content_parts = [str(row[col]) for col in text_columns if col in row and row[col]]
                    if content_parts:
                        content = " ".join(content_parts)
                        metadata = {k: v for k, v in row.items() if k not in text_columns}
                        documents.append({'id': f"csv_row_{i}", 'content': content, 'metadata': metadata})
            return documents
        except Exception as e:
            raise FileOperationError(f"Ошибка чтения CSV {filepath}: {e}") from e

    def load_from_json(self, filepath: str, content_field: str, id_field: Optional[str] = None) -> List[Dict]:
        path = Path(filepath)
        is_valid, reason = validate_path(
            path, allowed_dir=self.allowed_data_dir, allowed_exts={".json"}, max_size_mb=50
        )
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            documents = []
            if isinstance(data, list):
                for i, item in enumerate(data):
                    if content_field in item:
                        doc_id = item.get(id_field, f"json_item_{i}") if id_field else f"json_item_{i}"
                        content = str(item[content_field])
                        metadata = {k: v for k, v in item.items() if k not in [content_field, id_field]}
                        documents.append({'id': doc_id, 'content': content, 'metadata': metadata})
            return documents
        except Exception as e:
            raise FileOperationError(f"Ошибка чтения JSON {filepath}: {e}") from e

    def load_from_url(self, url: str, timeout: int = 30) -> str:
        try:
            response = requests.get(url, timeout=timeout)
            response.raise_for_status()
            if 'text/html' in response.headers.get('content-type', ''):
                if BS4_AVAILABLE:
                    soup = BeautifulSoup(response.text, 'html.parser')
                    for script in soup(["script", "style"]):
                        script.decompose()
                    return soup.get_text()
                else:
                    return response.text
            else:
                return response.text
        except Exception as e:
            raise ProcessingError(f"Ошибка загрузки URL {url}: {e}") from e

    async def load_from_urls_async(self, urls: List[str], timeout: int = 30, concurrency_limit: int = 10) -> List[Dict]:
        semaphore = asyncio.Semaphore(concurrency_limit)
        async def fetch_url(session, url):
            try:
                async with session.get(url, timeout=timeout) as response:
                    content = await response.text()
                    if 'text/html' in response.headers.get('content-type', ''):
                        if BS4_AVAILABLE:
                            soup = BeautifulSoup(content, 'html.parser')
                            for script in soup(["script", "style"]):
                                script.decompose()
                            content = soup.get_text()
                    return {
                        'id': hashlib.md5(url.encode()).hexdigest(),
                        'content': content,
                        'metadata': {'source_url': url, 'status': 'success'}
                    }
            except Exception as e:
                return {
                    'id': hashlib.md5(url.encode()).hexdigest(),
                    'content': "",
                    'metadata': {'source_url': url, 'status': 'error', 'error': str(e)}
                }
        async with aiohttp.ClientSession() as session:
            tasks = [fetch_url(session, url) for url in urls]
            results = await asyncio.gather(*tasks)
            return [r for r in results if r['content']]

class RAGAnalytics:
    """Аналитика и мониторинг RAG системы"""
    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline
        self.query_log = []

    def log_query(self, query: str, results_count: int, processing_time: float):
        self.query_log.append({
            'timestamp': time.time(),
            'query': query,
            'results_count': results_count,
            'processing_time': processing_time
        })

    def get_query_stats(self) -> Dict:
        if not self.query_log:
            return {"message": "Нет данных по запросам"}
        processing_times = [log['processing_time'] for log in self.query_log]
        results_counts = [log['results_count'] for log in self.query_log]
        return {
            'total_queries': len(self.query_log),
            'avg_processing_time': sum(processing_times) / len(processing_times),
            'max_processing_time': max(processing_times),
            'min_processing_time': min(processing_times),
            'avg_results_count': sum(results_counts) / len(results_counts),
            'queries_per_hour': len([q for q in self.query_log if time.time() - q['timestamp'] < 3600])
        }

    def analyze_collection_content(self) -> Dict:
        try:
            all_data = self.rag.collection.get()
            documents = all_data.get('documents', [])
            metadatas = all_data.get('metadatas', [])
            if not documents:
                return {"message": "Коллекция пуста"}
            total_docs = len(documents)
            total_chars = sum(len(doc) for doc in documents)
            avg_doc_length = total_chars / total_docs
            categories = {}
            languages = {}
            for metadata in metadatas:
                if metadata:
                    if 'category' in metadata:
                        cat = metadata['category']
                        categories[cat] = categories.get(cat, 0) + 1
                    if 'language' in metadata:
                        lang = metadata['language']
                        languages[lang] = languages.get(lang, 0) + 1
            return {
                'total_documents': total_docs,
                'total_characters': total_chars,
                'average_document_length': avg_doc_length,
                'categories_distribution': categories,
                'languages_distribution': languages,
                'longest_document': max(len(doc) for doc in documents),
                'shortest_document': min(len(doc) for doc in documents)
            }
        except Exception as e:
            return {"error": f"Ошибка анализа: {e}"}

class RAGWebInterface:
    """Простой веб-интерфейс для RAG системы"""
    def __init__(self, rag_pipeline: AdvancedRAGPipeline, analytics: RAGAnalytics):
        self.rag = rag_pipeline
        self.analytics = analytics

    def generate_html_interface(self) -> str:
        html_template = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>RAG Pipeline Interface</title>
            <meta charset="UTF-8">
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; background-color: #f5f5f5; }
                .container { max-width: 1200px; margin: 0 auto; background: white; padding: 20px; border-radius: 8px; }
                .search-box { width: 70%; padding: 10px; font-size: 16px; border: 1px solid #ddd; border-radius: 4px; }
                .search-btn { padding: 10px 20px; background: #007bff; color: white; border: none; border-radius: 4px; cursor: pointer; }
                .result { margin: 15px 0; padding: 15px; border: 1px solid #ddd; border-radius: 4px; background: #f9f9f9; }
                .score { color: #666; font-size: 14px; }
                .metadata { color: #888; font-size: 12px; margin-top: 5px; }
                .stats { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px; margin: 20px 0; }
                .stat-card { padding: 15px; background: #e9ecef; border-radius: 4px; text-align: center; }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>🔍 RAG Pipeline Interface</h1>
                <div class="search-section">
                    <input type="text" class="search-box" id="searchInput" placeholder="Введите ваш запрос...">
                    <button class="search-btn" onclick="performSearch()">Поиск</button>
                </div>
                <div id="results"></div>
                <h2>📊 Статистика системы</h2>
                <div class="stats" id="stats">
                    <div class="stat-card">
                        <h3>Запросы</h3>
                        <div id="queryCount">Загрузка...</div>
                    </div>
                    <div class="stat-card">
                        <h3>Среднее время</h3>
                        <div id="avgTime">Загрузка...</div>
                    </div>
                </div>
                <h2>📈 Управление данными</h2>
                <div style="margin: 20px 0;">
                    <button onclick="exportData()" style="margin-right: 10px; padding: 8px 16px; background: #28a745; color: white; border: none; border-radius: 4px;">Экспорт данных</button>
                    <button onclick="clearCollection()" style="padding: 8px 16px; background: #dc3545; color: white; border: none; border-radius: 4px;">Очистить коллекцию</button>
                </div>
            </div>
            <script>
                async function performSearch() {
                    const query = document.getElementById('searchInput').value;
                    if (!query.trim()) return;
                    const resultsDiv = document.getElementById('results');
                    resultsDiv.innerHTML = '<div>Поиск...</div>';
                    setTimeout(() => {
                        resultsDiv.innerHTML = `
                            <h3>Результаты поиска для: "${query}"</h3>
                            <div class="result">
                                <strong>Документ 1</strong>
                                <div class="score">Релевантность: 0.85</div>
                                <p>Пример найденного контента...</p>
                                <div class="metadata">Метаданные: category=ai, timestamp=2024-01-01</div>
                            </div>
                        `;
                    }, 1000);
                }
                function loadStats() {
                    document.getElementById('docCount').textContent = '150';
                    document.getElementById('queryCount').textContent = '45';
                    document.getElementById('avgTime').textContent = '0.3s';
                }
                function exportData() {
                    alert('Функция экспорта будет реализована на сервере');
                }
                function clearCollection() {
                    if (confirm('Вы уверены, что хотите очистить коллекцию?')) {
                        alert('Функция очистки будет реализована на сервере');
                    }
                }
                window.onload = loadStats;
                document.getElementById('searchInput').addEventListener('keypress', function(e) {
                    if (e.key === 'Enter') {
                        performSearch();
                    }
                });
            </script>
        </body>
        </html>
        """
        return html_template

    def save_interface(self, filepath: str = "rag_interface.html"):
        html_content = self.generate_html_interface()
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print(f"HTML интерфейс сохранен в {filepath}")

class RAGBenchmarking:
    """Бенчмаркинг и тестирование RAG системы"""
    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline

    def create_test_dataset(self, size: int = 100) -> List[Document]:
        test_topics = [
            "машинное обучение", "искусственный интеллект", "нейронные сети",
            "векторные базы данных", "обработка естественного языка",
            "компьютерное зрение", "глубокое обучение", "анализ данных",
            "программирование на Python", "веб-разработка"
        ]
        documents = []
        for i in range(size):
            topic = test_topics[i % len(test_topics)]
            content = self._generate_synthetic_content(topic, i)
            doc = Document(
                id=f"test_doc_{i}",
                content=content,
                metadata={
                    "topic": topic,
                    "test_id": i,
                    "synthetic": True
                }
            )
            documents.append(doc)
        return documents

    def _generate_synthetic_content(self, topic: str, doc_id: int) -> str:
        templates = [
            f"{topic} является важной областью современной технологии. Документ номер {doc_id} содержит подробную информацию о применении {topic} в различных сферах.",
            f"В этом документе рассматриваются основные принципы {topic}. Это руководство номер {doc_id} поможет понять ключевые концепции.",
            f"Практическое применение {topic} демонстрируется в примере {doc_id}. Здесь представлены лучшие практики и методы."
        ]
        base_content = templates[doc_id % len(templates)]
        additional = f" Дополнительные детали включают технические аспекты, примеры использования и рекомендации экспертов в области {topic}."
        return base_content + additional

    def benchmark_search_performance(self, queries: List[str], n_runs: int = 10) -> Dict:
        results = {
            'queries': [],
            'avg_time': 0,
            'total_time': 0,
            'fastest_query': None,
            'slowest_query': None
        }
        total_time = 0
        fastest_time = float('inf')
        slowest_time = 0
        for query in queries:
            query_times = []
            for _ in range(n_runs):
                start_time = time.time()
                self.rag.search(query, n_results=5)
                end_time = time.time()
                query_time = end_time - start_time
                query_times.append(query_time)
                total_time += query_time
            avg_query_time = sum(query_times) / len(query_times)
            min_query_time = min(query_times)
            max_query_time = max(query_times)
            if min_query_time < fastest_time:
                fastest_time = min_query_time
                results['fastest_query'] = query
            if max_query_time > slowest_time:
                slowest_time = max_query_time
                results['slowest_query'] = query
            results['queries'].append({
                'query': query,
                'avg_time': avg_query_time,
                'min_time': min_query_time,
                'max_time': max_query_time,
                'runs': n_runs
            })
        results['avg_time'] = total_time / (len(queries) * n_runs)
        results['total_time'] = total_time
        results['fastest_time'] = fastest_time
        results['slowest_time'] = slowest_time
        return results

    def evaluate_retrieval_quality(self, test_queries: List[Dict]) -> Dict:
        metrics = {
            'precision_at_k': [],
            'recall_at_k': [],
            'average_precision': []
        }
        for test_case in test_queries:
            query = test_case['query']
            relevant_ids = set(test_case['relevant_doc_ids'])
            results = self.rag.search(query, n_results=10)
            retrieved_ids = [r.document_id for r in results]
            for k in [1, 3, 5, 10]:
                if k <= len(retrieved_ids):
                    retrieved_k = set(retrieved_ids[:k])
                    precision = len(retrieved_k & relevant_ids) / k
                    recall = len(retrieved_k & relevant_ids) / len(relevant_ids) if relevant_ids else 0
                    metrics['precision_at_k'].append(precision)
                    metrics['recall_at_k'].append(recall)
            ap = self._calculate_average_precision(retrieved_ids, relevant_ids)
            metrics['average_precision'].append(ap)
        final_metrics = {
            'mean_precision_at_k': sum(metrics['precision_at_k']) / len(metrics['precision_at_k']) if metrics['precision_at_k'] else 0,
            'mean_recall_at_k': sum(metrics['recall_at_k']) / len(metrics['recall_at_k']) if metrics['recall_at_k'] else 0,
            'mean_average_precision': sum(metrics['average_precision']) / len(metrics['average_precision']) if metrics['average_precision'] else 0
        }
        return final_metrics

    def _calculate_average_precision(self, retrieved_ids: List[str], relevant_ids: set) -> float:
        if not relevant_ids:
            return 0.0
        precisions = []
        relevant_count = 0
        for i, doc_id in enumerate(retrieved_ids):
            if doc_id in relevant_ids:
                relevant_count += 1
                precision = relevant_count / (i + 1)
                precisions.append(precision)
        return sum(precisions) / len(relevant_ids) if precisions else 0.0
